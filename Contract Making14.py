from docx import*
from tkinter import *
import webbrowser
from tkinter.scrolledtext import ScrolledText
from tkinter import ttk
from tkinter import messagebox
from tkcalendar import Calendar
import datetime
from docx import Document
from docx import *
from docxcompose.composer import Composer
from docxcompose.composer import *
from docxcompose import *
from docx import Document
from tkcalendar import DateEntry
from docx.shared import Inches
import calendar
from random import *
from math import trunc




#doc_date_docx=Document ()
##composer = Composer(contract)


window = Tk()
window.geometry('1350x700')
window.resizable(0,0)
window.title('Contract Making')
window.config(background='white')


############################################################################# Get Input
def get_data():
    i1=doc_date_entry.get_date ()
    i2=combo_box_date_style.get ()
    i3=combo_box_ref_style.get ()
    i4=permeable_style=combo_box_start.get ()
    i5=combo_box_p_term.get ()
    i6=combo_box_d_term.get ()
    i7=combo_box_customer_origin.get ()
    i8=combo_box_activity.get ()
    i9=customer_name=customer_entry.get ()
    i10=customer_address= entry_customer_address.get()
    i11=our_comapny=combo_box_company.get ()
    i12=combo_box_trasaction.get ()
    i13=combo_box_currency.get ()
    i14=amount.entry.get ()
    i15=combo_box_gauranty.get ()
    i16=combo_box_sign.get ()
    i17=combo_box_font.get ()
    i18=combo_box_size.get ()
    i19=combo_box_templete.get ()
    i20=combo_box_border.get ()
    i21=combo_box_force.get ()
    i22=combo_box_general.get ()
    
    



############################################################################# Action on Input

#Style Source of Document

    if i19=='1':
        contract=Document ('style 1.docx')
    elif i19=='2':
        contract=Document ('style 2.docx')
        
    if i14=='':
        i14='0'

    doc_date=datetime.datetime.strftime(i1,'%b %d, %Y')


    ln=1
    abr=''
    for l in i9:
        if ln<4:
            abr=abr+l
            ln=ln+1
    

    if i2=='January 01, 2022':
        only_date=datetime.datetime.strftime(i1,'%b %d, %Y')
        p1=doc_date='Date: '+ datetime.datetime.strftime(i1,'%b %d, %Y')

    if i3=='ABC2201-01-1':
        m=0
        short_date=''
        for ii3 in datetime.datetime.strftime(i1,'%Y%m-%d'):
            if m>1:
                short_date=short_date+ii3
            m=m+1

         
        p2='Ref. No.: '+ abr+short_date+'-'+str(randint(1,9))


        

    our_company_name_1= '1 CO., LIMITED'
    our_company_name_address='FLAT/RM A 12/F ZJ 300, 300 LOCKHART ROAD, WAN CHAI, HONG KONG'



    
    
    if i11=='1':
        if i4=="1":
            p3='This sales contract (hereinafter referred to as the “contract”) is entered into on ' + only_date + '(the “effective date”), by and between ' + i9 + ' with an address of ' + i10 + ' (Hereinafter referred to as the “seller) and ' + our_company_name_1 + ' with an address of ' + our_company_name_address + ' (hereinafter referred to as the “customer”) (collectively referred to as the “parties”).'
    



    p4='The seller hereby agrees to sell the goods to the customer for an amount of ' + i13+' ' + i14 + ' by ' + i5 +'.'
        
    full_product=[['ETHOXYLATED FATTY ALCOHOL', 3154.00,'38231900' ,'Pt-Co: Less than 40','Hydroxyl Value: 89~99 mgKOH/g','C:1214','Packing: 190 Kg Drum'] , ['PVC ACRYLIC PROCESSING AID',1977.00 , '39069090','Purity: 99.5','Plastic Auxiliary Agents','Granularity (40 mesh): 98 Min.', 'Packing: 25 Kg Bag' ], ['SILICON DIOXIDE', 1250.00, '28112200', 'CAS No. 14808-60-7', 'Purity: 99%', 'Density: 2.6 g/mL', 'Packing: 25kg/Bag, 22MT/20 FT']]
    full_pod=[['Malaysia','Port Kalng', 'Penang Port'],['Singapore', 'Singapore Port'], ['UAE', 'Jebel Ali Port'], ['Oman','Sohar Port']]
    if i7==  'China':
        pol='Hong Kong Port, Hong Kong'
        if i8=='Chemical':
            full_strategy=[[full_product[0],1910000,0,0,0] ,[full_product[1],380000,7800,9630,0],[full_product[2],37600,46200,0,0]]
            ifull_strategy=randint(0,len(full_strategy)-1)
            strategy_group=full_strategy[ifull_strategy]
            final_strategy_group=[]
            for istrategy_group in range (len (strategy_group)):
                if istrategy_group>0:
                    if strategy_group [istrategy_group]!=0 and strategy_group [istrategy_group] > float (i14):
                        final_strategy_group=final_strategy_group+[istrategy_group]

            ifinal_strategy_group=randint(0,len(final_strategy_group)-1)
            final_strategy=final_strategy_group[ifinal_strategy_group]
            pod_group= full_pod[final_strategy_group[ifinal_strategy_group]]
            ipod_group=randint(1,len(pod_group)-1)
            pod=str (pod_group[ipod_group])+', '+str (pod_group[0])
            product_raw=strategy_group [0]
            
            
                 



            
            p5='The delivery of the goods (hereinafter referred to as the “delivery”) will by CFR '+ pod + ' '+  i6 + '\n' + 'The loading of the goods (hereinafter referred to as the “loading”) will be at ' + pol + '.'

    if i15=='Refund+No Gauranty':
        p6= 'Productss delivered not as described to the Buyer may be refunded for up to total eligible refund amount here: ' + i13 +' ' +i14 +'.Neither party makes any representations or grants any warranties, express or implied, either in fact or by operation of law, by statute or otherwise, and each party specifically disclaims any other warranties.'
        
    if i22=='1':
        p7='The risk of loss or damage for the goods will be on the seller until the products pass upon delivery to the buyer or its designee.'
        p8='Under no circumstances will the seller be liable for any indirect, special, consequential, or punitive damages (including lost profits) arising out of or relating to this agreement or the transactions it contemplates (whether for breach of contract, tort, negligence, or other form of action).'
        p9='In the event that any provision of this agreement is found to be void and unenforceable by a court of competent jurisdiction, then the remaining provisions will remain in force in accordance with the parties’ intention.'
        p10='This agreement contains the entire agreement and understanding among the parties hereto with respect to the subject matter hereof, and supersedes all prior agreements, understandings, inducements and conditions, express or implied, oral or written, of any nature whatsoever with respect to the subject matter hereof. The express terms hereof control and supersede any course of performance and/or usage of the trade inconsistent with any of the terms hereof.'
    
    if i16=='1' and i11=='1' :
        p11='The parties hereby agree to the terms and conditions set forth in this agreement and such is demonstrated throughout their signatures below:'+'\n'+'\n'+'Seller: '+ i9+'\n'+ p1+ '\n'+'\n'+ 'Buyer: '+our_company_name_1+'\n'+ p1



#Adding Sales Agreement Title (Heading)

        contract.add_heading('Sales Agreement', 1)

        contract.add_paragraph(p1+'\n')
        
        contract.add_paragraph(p2+'\n'+'\n')
        
        contract.add_heading('Parties', 2)

        contract.add_paragraph(p3+'\n'+'\n')

        contract.add_heading('Product', 2)

        
##        co1=p1+'\n'+p2+'\n'+'\n'+p3+'\n'+'\n'+'Product'+'\n'
##        
##        contract.add_paragraph(co1)

#Size of Table
        
        table=contract.add_table(rows = 4, cols = 4)

#Style of Table

        table.style ='table 1'


        
        row = table.rows[0].cells
        row[0].text = 'Description'
        row[1].text = 'QTY'
        row[2].text= 'Unit Price ' + i13
        row[3].text= 'Total ' + i13
        row1 = table.rows[1].cells
        row1[0].text = str(strategy_group [0][0])+'\n'+str(strategy_group [0][3])+'\n'+str(strategy_group [0][4])+'\n'+str(strategy_group [0][5])+'\n'+str(strategy_group [0][6])+'\n'+'HS CODE='+str(strategy_group [0][2])

# round down to two descimal
        
        def round2 (f):
           return((trunc(100*f))/100)

# choosing the number of products (rows)

        if float (i14)<100000:
            no_product=1
        elif float (i14):
            no_product=2
        else:
            no_product=3

# finding right unit price, QTY and discount
                        
        price_factor= (round2((uniform(0,2)-1)/10))+1
        unit_price_without_discount_1=float(strategy_group [0][1])*price_factor
        qty= trunc (float (i14) /unit_price_without_discount_1)
        unit_price_without_discount_2= round2 (float (i14) /qty)
        unit_price_with_discount_1=unit_price_without_discount_2+ round2(randint(100,300)/qty)
        total_line=round2(qty*unit_price_with_discount_1)
        discount=round2(total_line-float(i14))

# filling table with the numbers
        
        row1[2].text = str(unit_price_with_discount_1)     
        row1[1].text = str (qty)
        row1[3].text = str(total_line)
        row2= table.rows[2].cells
        row2[2].text='Discount '+ i13
        row2[3].text=str(discount)
        row3= table.rows[3].cells
        row3[2].text='Total Amount '+ i13
        row3[3].text=i14


        contract.add_heading('Price and payments', 2)

        contract.add_paragraph(p4+'\n'+'\n')




        contract.add_heading('Delivery and shipping', 2)

        contract.add_paragraph(p5+'\n'+'\n')

        

        contract.add_heading('Warranties', 2)

        contract.add_paragraph(p6+'\n'+'\n')



        contract.add_heading('Risk of loss and title', 2)

        contract.add_paragraph(p7+'\n'+'\n')


        
        contract.add_heading('Limitation of liability', 2)

        contract.add_paragraph(p8+'\n'+'\n')



        contract.add_heading('Severability', 2)

        contract.add_paragraph(p9+'\n'+'\n')



        contract.add_heading('Entire agreement', 2)

        contract.add_paragraph(p10+'\n'+'\n')        




        contract.add_heading('Signature and Date', 2)

        contract.add_paragraph(p11+'\n'+'\n')        




        
##        
##        co2='\n'+'\n'+p4+'\n'+'\n'+p5+'\n'+'\n'+p6+'\n'+'\n'+p7+'\n'+'\n'+p8
##        contract.add_paragraph(co2)    

# Saving the Contract 
        contract.save('contract11.docx')




##        total_line_2="{:.2f}".format (total_line)





############################################################################# USER INTERFACE

Button(window, text = "SUBMIT",
       
       font=("calibri 14 bold"),
	command = get_data).place(x=1100,y=600)






lbl_doc_date= Label(window,
            text='DOCUMENT DATE',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_doc_date.place(x=10,y=10)
doc_date_entry = DateEntry(window, locale='en_US', date_pattern='y-mm-dd')
doc_date_entry.place(x=150,y=10)


############################################################################# LINE 1
lbl_date_style= Label(window,
            text='DATE STYLE',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_date_style.place(x=250,y=10)


combo_box_date_style=ttk.Combobox (window,
                                 width=15
                                 )
combo_box_date_style.place(x=340,y=10)


combo_box_date_style['values']=['2022-01-01' ,'January 01, 2022', '01 JAN 2022']

combo_box_date_style.current(1)


#############################################################################
lbl_ref_style= Label(window,
            text='REF NUMBER STYLE',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_ref_style.place(x=460,y=10)


combo_box_ref_style=ttk.Combobox (window,
                                 width=15
                                 )
combo_box_ref_style.place(x=610,y=10)


combo_box_ref_style['values']=['ABC2201-01-1' ,'ABC0122-01-1', 'ABC22-0101-1']

combo_box_ref_style.current(0)




#############################################################################

lbl_start= Label(window,
            text='PERMEABLE STYLE',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_start.place(x=730,y=10)


combo_box_start=ttk.Combobox (window,
                                 width=15
                                 )
combo_box_start.place(x=870,y=10)


combo_box_start['values']=['1','2']

combo_box_start.current(0)

############################################################################# LINE 2

lbl_p_term= Label(window,
            text='PAYMENT TERM',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_p_term.place(x=10,y=60)


combo_box_p_term=ttk.Combobox (window,
                                 width=15
                                 )
combo_box_p_term.place(x=150,y=60)


combo_box_p_term['values']=['100% In Advance','100% Against BL Copy', '100% Against BL Releasing', 'CIA (Cash In Advance','Buy Now,Pay Later, with 5% Interest Rate' ]

combo_box_p_term.current(0)

#############################################################################

lbl_d_term= Label(window,
            text='DELIVERY TIME',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_d_term.place(x=270,y=60)


combo_box_d_term=ttk.Combobox (window,
                                 width=50
                                 )
combo_box_d_term.place(x=385,y=60)
d_term_i=datetime.datetime. now ()
d_term_year=str(d_term_i.year)
d_term_v=['Within two months after the full payment','Eight to Twelve weeks after the full payment' ]
for v in range(1,13):

        d_term_v.append('in ' +calendar.month_name[v]+ ' '+d_term_year+ ' ,Subject to Payment in 30 Days')
        

                        
                            


combo_box_d_term['values']=d_term_v

combo_box_d_term.current(0)


#############################################################################

##lbl_pol= Label(window,
##            text='POL',
##            font=("calibri 12 bold"),
##            background=('White'),
##
##            )
##lbl_pol.place(x=720,y=60)
##
##
##combo_box_pol=ttk.Combobox (window,
##                                 width=30
##                                 )
##combo_box_pol.place(x=760,y=60)
##
##
##combo_box_pol['values']=['Mundra Port, India','Busan Port, South Korea', 'Kaohsiung Port, Taiwan', 'Laem Chabang Port, Thailand', 'Mersin Port, Turkey','Hong Kong Port, Hong Kong',
##                         'Port Klang, Malaysia','Penang Port, Malaysia' ,'Singapore Port, Singapore', 'Jebel Ali Port, UAE', 'Sohar Port, Oman','Jakarta Port, Indonesia',
##                         'Hamburk Port, Germany', 'Felixstowe Port, UK']
##
##combo_box_pol.current(0)
###############################################################################
##
##lbl_pod= Label(window,
##            text='POD',
##            font=("calibri 12 bold"),
##            background=('White'),
##
##            )
##lbl_pod.place(x=970,y=60)
##
##
##combo_box_pod=ttk.Combobox (window,
##                                 width=30
##                                 )
##combo_box_pod.place(x=1010,y=60)
##
##
##combo_box_pod['values']=['Port Klang, Malaysia','Penang Port, Malaysia' ,'Singapore Port, Singapore', 'Jebel Ali Port, UAE', 'Sohar Port, Oman','Jakarta Port, Indonesia']
##
##combo_box_pod.current(0)


#############################################################################

lbl_customer_origin= Label(window,
            text='CUSTOMER ORIGIN',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_customer_origin.place(x=720,y=60)


combo_box_customer_origin=ttk.Combobox (window,
                                 width=15
                                 )
combo_box_customer_origin.place(x=870,y=60)


combo_box_customer_origin['values']=['India','South Korea', 'Taiwan', 'Thailand', 'Turkey','Hong Kong',
                         'Malaysia','Singapore', 'UAE', 'Oman','Indonesia',
                         'Germany', 'UK', 'Kyrgyzstan', 'China']

combo_box_customer_origin.current(14)
#############################################################################

lbl_activity= Label(window,
            text='ACTIVITY',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_activity.place(x=990,y=60)


combo_box_activity=ttk.Combobox (window,
                                 width=30
                                 )
combo_box_activity.place(x=1070,y=60)


combo_box_activity['values']=['Chemical', 'Petrochemical','Petroleum Product', 'Crude Oil' ,'Valve', 'Industrial Pump/Compressor', 'Refrigerator Comperssor', 'AC Comperssor','CAR', 'Motrocycle', 'Bicycle', 'Food', 'Textile', 'Metalwork', 'SMT', 'Steel', 'Aluminium', 'Packing Machine', 'Architecture', 'Computer', 'TV', 'Home Appliances', 'Construction Material', 'Diesel Engine']

combo_box_activity.current(0)


############################################################################# LINE 3

customer = Label(window,
            text='CUSTOMER NAME',
            font=("calibri 12 bold"),
            background=('White'),
            )
customer.place(x=10,y=110)
customer_entry = Entry(window,width=40, background=('light grey'),font=("calibri 15 italic"))
customer_entry.place(x=150,y=110)


#############################################################################

lbl_customer_address= Label(window,
            text='ADDRESS',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_customer_address.place(x=570,y=110)

entry_customer_address= Entry(window,
                          width=65,
                          background=('light grey'),
                          font=("calibri 15 italic")
                          )
entry_customer_address.place(x=650,y=110)


############################################################################# LINE 4
lbl_company= Label(window,
            text='OUR COMPANY',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_company.place(x=10,y=160)

combo_box_company=ttk.Combobox (window,
                                 width=15
                                 )
combo_box_company.place(x=150,y=160)


combo_box_company['values']=['1','2','3','4','5','6','7','8']

combo_box_company.current(0)



#############################################################################
lbl_trasaction= Label(window,
            text='OUT/IN',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_trasaction.place(x=270,y=160)

combo_box_trasaction=ttk.Combobox (window,
                                   width=15
                                   )
combo_box_trasaction.place(x=340,y=160)


combo_box_trasaction['values']=['OUTWARD','INWARD']

combo_box_trasaction.current(0)

#############################################################################
lbl_currency= Label(window,
            text='CURRENCY',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_currency.place(x=460,y=160)


combo_box_currency=ttk.Combobox (window,
                                 width=15
                                 )
combo_box_currency.place(x=550,y=160)


combo_box_currency['values']=['USD','EURO','CNY']

combo_box_currency.current(0)

#############################################################################
amount = Label(window,
            text='AMOUNT',
            font=("calibri 12 bold"),
            background=('White'),
            )
amount.place(x=670,y=160)
amount.entry = Entry(window,width=30, background=('light grey'),font=("calibri 15 italic"))
amount.entry.place(x=750,y=160)




############################################################################# LINE 5
lbl_gauranty= Label(window,
            text='GAURANTY',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_gauranty.place(x=10,y=210)


combo_box_gauranty=ttk.Combobox (window,
                                 width=15
                                 )
combo_box_gauranty.place(x=150,y=210)


combo_box_gauranty['values']=['12 Months' ,'24 Months', 'As is+Inspection','Refund+No Gauranty']

combo_box_gauranty.current(3)
#############################################################################
lbl_sign= Label(window,
            text='EXECUTION STYLE',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_sign.place(x=270,y=210)

combo_box_sign=ttk.Combobox (window,
                                   width=15
                                   )
combo_box_sign.place(x=400,y=210)


combo_box_sign['values']=['1','2']

combo_box_sign.current(0)

#############################################################################
lbl_font= Label(window,
            text='FONT',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_font.place(x=520,y=210)

combo_box_font=ttk.Combobox (window,
                                   width=15
                                   )
combo_box_font.place(x=570,y=210)


combo_box_font['values']=['Calibri','Times', 'Arial', 'Bookman']

combo_box_font.current(0)
#############################################################################
lbl_size= Label(window,
            text='FONT SIZE',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_size.place(x=690,y=210)

combo_box_size=ttk.Combobox (window,
                                   width=15
                                   )
combo_box_size.place(x=770,y=210)


combo_box_size['values']=['10','11', '12']

combo_box_size.current(1)

#############################################################################

lbl_templete= Label(window,
            text='TEMPLETE',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_templete.place(x=890,y=210)

combo_box_templete=ttk.Combobox (window,
                                   width=15
                                   )
combo_box_templete.place(x=980,y=210)


combo_box_templete['values']=['1','2', '3']

combo_box_templete.current(0)

#############################################################################
lbl_border= Label(window,
            text='BORDER',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_border.place(x=1100,y=210)

combo_box_border=ttk.Combobox (window,
                                   width=15
                                   )
combo_box_border.place(x=1170,y=210)


combo_box_border['values']=['Without','1', '2']

combo_box_border.current(0)

#############################################################################

lbl_force= Label(window,
            text='FORCE MAJEURE',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_force.place(x=10,y=260)

combo_box_force=ttk.Combobox (window,
                                   width=15
                                   )
combo_box_force.place(x=150,y=260)


combo_box_force['values']=['1','2', '3']

combo_box_force.current(0)
#############################################################################
lbl_general= Label(window,
            text='GENERAL PROVISIONS',
            font=("calibri 12 bold"),
            background=('White'),

            )
lbl_general.place(x=270,y=260)

combo_box_general=ttk.Combobox (window,
                                   width=15
                                   )
combo_box_general.place(x=440,y=260)


combo_box_general['values']=['1','2', '3']

combo_box_general.current(0)
#############################################################################






#############################################################################


window.mainloop()


